import json
import os
import time
import uuid
import threading
from pathlib import Path
from typing import Dict, Any

from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

import pypdf
import docx
import pandas as pd

class SecureParser:
    """Handles secure extraction of supported file formats, stripping active elements."""
    
    @staticmethod
    def process_pdf(file_path: Path) -> str:
        """Extracts text from PDF while natively ignoring Javascript and active forms."""
        text = []
        with open(file_path, "rb") as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        return "\n".join(text)

    @staticmethod
    def process_docx(file_path: Path) -> str:
        """Extracts plain text from Word documents ignoring macros and OLE objects."""
        doc = docx.Document(file_path)
        text = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text.append(cell.text)
        return "\n".join(text)

    @staticmethod
    def process_csv_excel(file_path: Path) -> str:
        """Parses spreadsheets, escapes formula injections (=+-@), and returns clean JSON."""
        ext = file_path.suffix.lower()
        if ext == ".csv":
            df = pd.read_csv(file_path, dtype=str).fillna("")
        else:
            df = pd.read_excel(file_path, dtype=str).fillna("")
            
        def escape_formula(val):
            if isinstance(val, str) and val and val[0] in ("=", "+", "-", "@"):
                return f"'{val}"
            return val
            
        # Escape cell values
        for col in df.columns:
            df[col] = df[col].apply(escape_formula)
            
        # Escape column headers if they are malicious
        new_cols = {}
        for col in df.columns:
            if isinstance(col, str) and col and col[0] in ("=", "+", "-", "@"):
                new_cols[col] = f"'{col}"
        if new_cols:
            df.rename(columns=new_cols, inplace=True)
            
        return df.to_json(orient="records", indent=2)

    @staticmethod
    def process_txt(file_path: Path) -> str:
        """Extracts raw text with resilient UTF-8 decoding."""
        try:
            return file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            return file_path.read_text(encoding="latin-1")

    @staticmethod
    def _verify_magic_bytes(file_path: Path, ext: str) -> bool:
        """Validates file signatures to prevent extension spoofing."""
        try:
            with open(file_path, "rb") as f:
                header = f.read(8)
                
            if ext == ".pdf":
                return header.startswith(b"%PDF")
            elif ext in (".docx", ".xlsx"):
                return header.startswith(b"PK\x03\x04")
            elif ext in (".csv", ".txt"):
                try:
                    header.decode("utf-8")
                    if header.startswith(b"MZ") or header.startswith(b"\x7fELF"):
                        return False
                    return True
                except UnicodeDecodeError:
                    try:
                        header.decode("latin-1")
                        if header.startswith(b"MZ") or header.startswith(b"\x7fELF"):
                            return False
                        return True
                    except Exception:
                        return False
            elif ext == ".xls":
                # Legacy Excel OLE format
                return header.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1")
            return True
        except Exception:
            return False


class QuarantineEventHandler(FileSystemEventHandler):
    def __init__(self, parser_service: 'QuarantineService'):
        self.parser_service = parser_service

    def on_created(self, event):
        if not event.is_directory:
            # Execute parsing asynchronously so watchdog listener remains responsive
            threading.Thread(target=self.parser_service.process_file, args=(Path(event.src_path),)).start()


class QuarantineService:
    def __init__(self, workspace_dir: str | Path, audit_log_path: str | Path):
        self.workspace_dir = Path(workspace_dir).resolve()
        self.quarantine_dir = self.workspace_dir / "quarantine"
        self.sanitized_dir = self.workspace_dir / "sanitized"
        self.audit_log_path = Path(audit_log_path).resolve()
        
        self.quarantine_dir.mkdir(parents=True, exist_ok=True)
        self.sanitized_dir.mkdir(parents=True, exist_ok=True)
        
        self.observer = Observer()
        self.handler = QuarantineEventHandler(self)
        self.observer.schedule(self.handler, str(self.quarantine_dir), recursive=False)

    def start(self):
        self.observer.start()
        
    def stop(self):
        self.observer.stop()
        self.observer.join()
        
    def _write_audit_log(self, entry: Dict[str, Any]) -> None:
        """Logs ingestion results to the central audit JSON."""
        logs = []
        try:
            if self.audit_log_path.exists():
                with open(self.audit_log_path, "r", encoding="utf-8") as f:
                    content = f.read().strip()
                    if content:
                        logs = json.loads(content)
        except Exception:
            logs = []

        logs.append(entry)
        try:
            with open(self.audit_log_path, "w", encoding="utf-8") as f:
                json.dump(logs, f, indent=2)
        except Exception:
            pass
            
    def _wait_for_file_lock(self, file_path: Path, timeout: int = 5) -> bool:
        """Ensures the downloading process has finished writing before ingestion."""
        start_time = time.time()
        while time.time() - start_time < timeout:
            try:
                # Rename checks if the file is exclusively locked by another process
                os.rename(file_path, file_path)
                return True
            except OSError:
                time.sleep(0.5)
        return False

    def process_file(self, file_path: Path):
        """Parses, sanitizes, and promotes a quarantined file."""
        start_time = time.time()
        file_name = file_path.name
        
        if not file_path.exists():
            return
            
        if not self._wait_for_file_lock(file_path):
            self._write_audit_log({
                "event_type": "quarantine_ingestion",
                "file_name": file_name,
                "mime_type": "unknown",
                "file_size_bytes": 0,
                "sanitization_status": "failed",
                "error": "File lock timeout - file might be stuck downloading.",
                "duration_ms": int((time.time() - start_time) * 1000),
                "extracted_characters_count": 0,
                "timestamp": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
            })
            return

        try:
            file_size = file_path.stat().st_size
            
            # Defensive Size Thresholding
            if file_size > 25 * 1024 * 1024:
                raise ValueError("size_limit_exceeded: File exceeds 25MB safety threshold.")
                
            ext = file_path.suffix.lower()
            
            # Magic Byte Signature Validation
            if not SecureParser._verify_magic_bytes(file_path, ext):
                raise ValueError(f"Magic byte signature mismatch for extension {ext}")
            
            output_content = ""
            mime_type = "unknown"
            
            # Secure Routing
            if ext == ".pdf":
                mime_type = "application/pdf"
                output_content = SecureParser.process_pdf(file_path)
            elif ext == ".docx":
                mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                output_content = SecureParser.process_docx(file_path)
            elif ext in (".csv", ".xlsx", ".xls"):
                mime_type = "text/csv" if ext == ".csv" else "application/vnd.ms-excel"
                output_content = SecureParser.process_csv_excel(file_path)
            elif ext == ".txt":
                mime_type = "text/plain"
                output_content = SecureParser.process_txt(file_path)
            else:
                raise ValueError(f"Unrecognized or hostile file format: {ext}")

            # Promote to safe workspace zone
            safe_id = str(uuid.uuid4())
            out_ext = ".json" if ext in (".csv", ".xlsx", ".xls") else ".txt"
            sanitized_path = self.sanitized_dir / f"{safe_id}{out_ext}"
            
            sanitized_path.write_text(output_content, encoding="utf-8")
            
            # Clean up the hostile payload
            file_path.unlink()
            
            status = "success"
            error_msg = None
            extracted_chars = len(output_content)
            
        except Exception as e:
            # Fail closed: file remains quarantined
            status = "quarantined_malicious" if isinstance(e, ValueError) else "failed"
            error_msg = str(e)
            extracted_chars = 0
            
        duration = int((time.time() - start_time) * 1000)
        self._write_audit_log({
            "event_type": "quarantine_ingestion",
            "file_name": file_name,
            "mime_type": mime_type if 'mime_type' in locals() else "unknown",
            "file_size_bytes": file_size if 'file_size' in locals() else 0,
            "sanitization_status": status,
            "error": error_msg,
            "duration_ms": duration,
            "extracted_characters_count": extracted_chars,
            "timestamp": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
        })
